# CMT114 Coursework
# student number:

import docx
import os, sys

def change_style(filepath, style='IEEE'):

	doc = docx.Document(filepath)

	'''
	思路是找到reference 然后通过reference 通过reference 判断 style 最后再修改in_text
	'''
	references = [ ]

	def REPLACE(x, old, new):

		old_len = len(old)

		while True:
			exist = False

			for i in range(len(x) - old_len + 1):

				sub_str = x[i:i + old_len]

				if sub_str == old:
					exist = True
					x = x[:i] + new + x[min(len(x), i + old_len):]

			if not exist:
				break

		return x

	for i in range(len(doc.paragraphs)):

		paragraph = doc.paragraphs[i]

		if "References" == paragraph.text:

			reference_start_index = i +  1

			#表明这一行后面都是reference

			break


	for j in range(reference_start_index,len(doc.paragraphs)):

		references.append(doc.paragraphs[j].text)

		#找到所有的reference

	if references[0][0] == "1":
		doc_style = "IEEE"

	else:
		doc_style = "APA"

	if doc_style ==  style:

		doc.save(filepath[:-5] + "_%s_style"%style+".docx")

		return 0
		#风格相同不做转换,直接保留

	if doc_style == "IEEE":


		#首先格式化references,找到每个reference的序号,前两个

		#下面将从IEEE风格转换为APA风格


		reference_dict = {

		}

		reference_list = []
		for reference in references:

			if REPLACE(reference," ","") == "":
				continue
			items = reference.split(",")

			date = items[-1]

			for j in range(len(items)):

				if len(items[j]) >= 20:

					#字符串长度超过20 认为是引用资料的书名

					break
			source =  ",".join(items[j: -1])

			authors = items[:j]

			index = authors[0].split(".")[0]

			authors[0] = ".".join(authors[0].split(".")[1:])

			year = ""

			for char in date:

				if str(char).isdigit():

					year += char

			authors  = list(map(lambda x : x.split(".")[-1] + "," + ".".join(x.split(".")[:-1]) + ".",authors))
			reference_dict[index] = {
				"year":year,
				"authors": authors,
				"source" :source,
			}

			reference_list.append((authors,source,year))

		#将reference_list 按 第一作者的字母排序

		reference_list = sorted(reference_list,key = lambda x : ord(x[0][0].split(",")[0][0]),reverse=False)

		#下面转换为APA格式的references

		for k in range(len(reference_list)):

			new_refence_data = reference_list[k]

			authors = new_refence_data[0]

			source  = new_refence_data[1]

			year  = new_refence_data[2]

			new_refence = ""

			print(authors)

			if len(authors)  == 1:

				new_refence += authors[0]

			else:
				new_refence += ",".join(authors[:-1])

				new_refence += ", and" + authors[-1]

			new_refence += " (%s)"%year + "."

			new_refence += source

			doc.paragraphs[reference_start_index + k].text =new_refence

		#下面转换为APA格式的in-text citation:

		reference_num = len(reference_list)

		for i in range(len(doc.paragraphs)):

			paragraph = doc.paragraphs[i]

			paragraph_text = paragraph.text

			for j in range(reference_num):

				if "[%s]"%str(j) in paragraph_text:

					#表明这段有in-text citation,首先将多个括号放在一起,比如[4],[5] 变为[4,5]

					paragraph_text = REPLACE(paragraph_text,"], [",",")

					new_paragraph_text = paragraph_text

					#下面搜索索引


					for m in range(len(paragraph_text)):

						if paragraph_text[m] == "[":

							for n in range(m + 1 , len(paragraph_text) -1):

								if paragraph_text[n] == "]":

										break

							citation_indexs = paragraph_text[m + 1:n].split(",")

							citation_replace = ""

							for citation_index in citation_indexs:

								citation_data = reference_dict[citation_index]

								names = list(map(lambda  x :REPLACE(x.split(",")[0]," ","") ,  citation_data["authors"]))

								if len(names) >= 3:

									citation_replace += names[0] + " et al."

								elif len(names) == 2:

									citation_replace += names[0] + " and " + names[1]

								else:

									citation_replace += names[0]

								citation_replace += ", " + citation_data["year"] +"; "

							citation_replace = "(" + citation_replace[:-2] + ")" #去掉最末尾的;

							origin_text = paragraph_text[m:n + 1]

							new_paragraph_text = REPLACE(new_paragraph_text,origin_text, citation_replace)


					doc.paragraphs[i].text = new_paragraph_text


		doc.save(filepath[:-5] + "_%s_style" % style + ".docx")

	if doc_style == "APA":

		#下面将从APA风格转换为IEEE风格
		reference_dict = {

		}

		reference_list = []

		reference_authors = []


		for reference in references:

			#先找到年份
			if REPLACE(reference," ","") == "":

				continue

			left_bracket = reference.index("(")

			right_bracket = reference.index(")")

			year = reference[left_bracket + 1:right_bracket]

			raw_authors = reference[:left_bracket].split(",")

			authors= []

			for k in range(int(len(raw_authors)/2)):
				authors.append(
					(REPLACE(REPLACE(raw_authors[0+2*k]," ",""),"and",""),REPLACE(raw_authors[1+2*k]," ",""))
				)

				reference_authors.append(REPLACE(REPLACE(raw_authors[0+2*k]," ",""),"and",""))

			#找到作者的名和姓

			source = reference[right_bracket + 1:]

			reference_list.append((year,authors,source))


		citation_index_assign = []
		#下面对文档里面的in-text进行引用赋予index
		for i in range(reference_start_index):
				paragraph = doc.paragraphs[i]

				paragraph_text = paragraph.text

				new_paragraph_text = paragraph_text


				for m in range(len(paragraph_text)):

					if paragraph_text[m] == "(":

						for n in range(m + 1, len(paragraph_text) - 1):

							if paragraph_text[n] == ")":

								break

						citations= REPLACE(paragraph_text[m + 1:n],"e.g.,","").split(";")

						if  citations == ['a person']:

							continue

						if  "http" in citations[0]:

							continue

						citation_data = []

						if str.isdigit(citations[0]):

							previous_content = paragraph_text[max(0,m-30):m].split(" ")

							for word in previous_content:

								word =  REPLACE(REPLACE(word,"^","")," ","")

								if word in reference_authors:

									citation_author = word

									break
							#print(previous_content)
							citation_year = citations[0]

							if citation_author + citation_year not in citation_index_assign:

								citation_index_assign.append(citation_author + citation_year)

							citation_data.append(citation_author + citation_year)

							#表明这个引用只有年份 需要找括号前找作者
						else:
							for citation in citations:

								citation_year = REPLACE(citation.split(",")[-1]," ","")

								if str.isdigit(citation_year) == False:

									continue
								citation_author = REPLACE(citation.split(",")[0],"et al.","")

								citation_author = REPLACE(citation_author.split("and")[0]," ","")

								citation_data.append(citation_author + citation_year)

								if citation_author + citation_year not in citation_index_assign:

									citation_index_assign.append(citation_author + citation_year)

						citation_indexs = []
						if citation_data == []:

							continue
						else:

							for citation in citation_data:

								citation_index = citation_index_assign.index(citation)

								citation_indexs.append(citation_index)

						citation_replace = "".join(["[%s]"%(index+1) for index in citation_indexs])

						origin_text = paragraph_text[m:n + 1]

						new_paragraph_text = REPLACE(new_paragraph_text,origin_text,citation_replace)


				doc.paragraphs[i].text = new_paragraph_text

		#下面转换为APA格式的references

		new_reference_dict = { }
		for k in range(len(reference_list)):

			new_refence_data = reference_list[k]

			authors = new_refence_data[1]

			source  = new_refence_data[2]

			year  = new_refence_data[0]

			new_refence = ""

			first_author = authors[0]

			citation_key = first_author[0] + year

			index = citation_index_assign.index(citation_key)  + 1

			new_refence += str(index) + "."

			author_names = [author[1] + author[0] for author in authors]

			for name in author_names:
				new_refence += name +", "

			new_refence += REPLACE(source,".","")

			new_refence += ", %s."%year

			new_reference_dict[index] = new_refence

		for k in range(len(reference_list)):

			doc.paragraphs[reference_start_index + k].text =new_reference_dict[k + 1]

		doc.save(filepath[:-5] + "_%s_style" % style + ".docx")


# ---- DO NOT CHANGE THE CODE BELOW ----
if __name__ == "__main__":
	change_style("IEEEexample_no_hyperlinks .docx",style="APA")
	#if len(sys.argv)<3: raise ValueError('Provide filename and style as input arguments')
	#filepath, style = sys.argv[1], sys.argv[2]
	#print('filepath is "{}"'.format(filepath))
	#print('target style is "{}"'.format(style))
	#change_style(filepath, style)
